#!/usr/bin/env python3
"""Build the bilingual ScientificParallax research report as polished DOCX files."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
PAPER_DIR = ROOT / "papers"
OUTPUT_DIR = ROOT / "artifacts" / "paper"
ASSET_DIR = OUTPUT_DIR / "assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203748"
MUTED = "5C6670"
LIGHT_BLUE = "EAF2F8"
LIGHT_GREEN = "E9F5EC"
LIGHT_AMBER = "FFF4E5"
LIGHT_RED = "FDECEC"
TABLE_FILL = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|https?://[^\s)]+)")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "CCD5DD", size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    """Apply fixed Word geometry: tblW, tblInd, tblGrid, and every tcW agree."""
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def mark_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(
    run,
    font: str,
    east_asia: str,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(
    paragraph, text: str, url: str, font: str, east_asia: str, size: float | None = None
) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_pr.extend([r_fonts, color, underline])
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        r_pr.append(sz)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, font: str, east_asia: str, size: float | None = None) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, font, east_asia, size=size)
        token = match.group(0)
        if token.startswith("http"):
            if "doi.org/" in token:
                label = "DOI"
            elif "arxiv.org/" in token:
                label = "arXiv"
            elif "proceedings.mlr.press/" in token:
                label = "PMLR"
            else:
                label = "Source"
            add_hyperlink(paragraph, label, token, font, east_asia, size=size)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, font, east_asia, size=size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, font, east_asia, size=size, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, font, east_asia, size=size)


def add_field_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def create_numbering(document: Document, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abs_ids = [
        int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abs_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "279")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, fmt, lvl_text, lvl_jc, p_pr])
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def configure_document(document: Document, language: str) -> tuple[str, str]:
    # Named override for the Chinese edition: LibreOffice on macOS does not
    # consistently honor w:eastAsia fallback when the primary font is Calibri.
    font = "Songti SC" if language == "zh" else "Calibri"
    east_asia = font if language == "zh" else "Arial"
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.element.rPr.rFonts.set(qn("w:ascii"), font)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), font)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    # A compact but readable leading keeps short paragraphs from being split
    # across pages by Word and headless LibreOffice in different places.
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.keep_together = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.element.rPr.rFonts.set(qn("w:ascii"), font)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), font)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = font
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.element.rPr.rFonts.set(qn("w:ascii"), font)
    caption.element.rPr.rFonts.set(qn("w:hAnsi"), font)
    caption.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = "SCIENTIFICPARALLAX"
    right = "证据触发的研究方向" if language == "zh" else "EVIDENCE-TRIGGERED RESEARCH DIRECTIONS"
    r1 = hp.add_run(left)
    set_run_font(r1, "Arial", east_asia, size=8.5, bold=True, color=MUTED)
    r2 = hp.add_run("\t" + right)
    set_run_font(r2, font if language == "zh" else "Arial", east_asia, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_field_page_number(fp)
    for run in fp.runs:
        set_run_font(run, "Arial", east_asia, size=8.5, color=MUTED)

    document.core_properties.author = "ScientificParallax Project"
    document.core_properties.subject = (
        "Evidence-triggered scientific direction generation by an LLM"
    )
    document.core_properties.comments = (
        "Generated from frozen local research artifacts; development report, not a novelty claim."
    )
    return font, east_asia


def add_cover(
    document: Document,
    title: str,
    subtitle: str,
    metadata: str,
    takeaway: str,
    language: str,
    font: str,
    east_asia: str,
) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker = "研究报告" if language == "zh" else "RESEARCH REPORT"
    run = p.add_run(kicker)
    set_run_font(
        run, font if language == "zh" else "Arial", east_asia, size=10, bold=True, color=BLUE
    )

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, font, east_asia, size=27 if language == "zh" else 26, bold=True, color=INK)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, font, east_asia, size=14, color=DARK_BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run(metadata)
    set_run_font(
        run, font if language == "zh" else "Arial", east_asia, size=9.5, italic=True, color=MUTED
    )

    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.45)
    p.paragraph_format.right_indent = Inches(0.45)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "18")
    left_border.set(qn("w:space"), "10")
    left_border.set(qn("w:color"), BLUE)
    borders.append(left_border)
    p_pr.append(borders)
    run = p.add_run(takeaway)
    set_run_font(run, font, east_asia, size=11, bold=True, color=INK)
    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Inches(8.5)
    body_section.page_height = Inches(11)
    body_section.top_margin = Inches(1)
    body_section.right_margin = Inches(1)
    body_section.bottom_margin = Inches(1)
    body_section.left_margin = Inches(1)
    body_section.header_distance = Inches(0.492)
    body_section.footer_distance = Inches(0.492)


def add_image(
    document: Document, path: Path, caption: str, alt: str, font: str, east_asia: str
) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(6.25))
    inline._inline.docPr.set("descr", alt)
    cp = document.add_paragraph(style="Caption")
    add_inline(cp, caption, font, east_asia, size=9)


def add_table(
    document: Document, rows: list[list[str]], table_index: int, font: str, east_asia: str
) -> None:
    if table_index == 1:
        widths = [1500, 2350, 2100, 3410]
    elif table_index == 2:
        widths = [1800, 850, 3270, 3440]
    elif table_index == 3:
        widths = [2100, 1500, 5760]
    else:
        cols = len(rows[0])
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, font, east_asia, size=9.2)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(INK)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def extract_front_matter(lines: list[str]) -> tuple[str, str, str, str, int]:
    nonempty = [(idx, line.strip()) for idx, line in enumerate(lines) if line.strip()]
    title_idx, title_line = nonempty[0]
    subtitle_idx, subtitle_line = nonempty[1]
    metadata_idx, metadata = nonempty[2]
    takeaway_idx, takeaway_line = nonempty[3]
    assert title_line.startswith("# ")
    assert subtitle_line.startswith("## ")
    assert takeaway_line.startswith("> ")
    return (
        title_line[2:].strip(),
        subtitle_line[3:].strip(),
        metadata,
        takeaway_line[2:].strip(),
        takeaway_idx + 1,
    )


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[str] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw.append(lines[idx].strip())
        idx += 1
    rows: list[list[str]] = []
    for pos, line in enumerate(raw):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if pos == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def render_markdown(
    document: Document, source: Path, language: str, image_map: dict[str, Path]
) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    title, subtitle, metadata, takeaway, idx = extract_front_matter(lines)
    font, east_asia = configure_document(document, language)
    document.core_properties.title = title
    add_cover(document, title, subtitle, metadata, takeaway, language, font, east_asia)
    bullet_id = create_numbering(document, "bullet")
    decimal_id = create_numbering(document, "decimal")
    table_index = 0
    explainer_paragraphs_remaining = 0
    compact_references = False
    compact_appendix = False

    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue

        if line.startswith("!["):
            match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
            if not match:
                raise ValueError(f"Malformed image marker: {line}")
            caption = match.group(1)
            marker = Path(match.group(2)).name
            add_image(document, image_map[marker], caption, caption, font, east_asia)
            idx += 1
            continue

        if line.startswith("|"):
            rows, idx = parse_table(lines, idx)
            table_index += 1
            add_table(document, rows, table_index, font, east_asia)
            continue

        if line.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline(p, line[4:].strip(), font, east_asia, size=13)
            idx += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            compact_references = heading in {"参考文献", "References"}
            compact_appendix = heading in {
                "附录：预先判据的实际结果",
                "Appendix: observed outcomes under the frozen rules",
            }
            if heading in {"1. 研究问题", "1. Research question"}:
                spacer = document.add_paragraph()
                spacer.paragraph_format.line_spacing = Pt(24)
                spacer.paragraph_format.space_after = Pt(6)
                spacer.paragraph_format.keep_with_next = True
                spacer_run = spacer.add_run("\u00a0")
                set_run_font(spacer_run, font, east_asia, size=1, color=WHITE)
            p = document.add_paragraph(style="Heading 1")
            add_inline(p, heading, font, east_asia, size=16)
            if heading in {
                "给不同背景读者的简短说明",
                "A short explanation for readers from any field",
            }:
                explainer_paragraphs_remaining = 2
            idx += 1
            continue

        if line.startswith("- "):
            p = document.add_paragraph()
            apply_numbering(p, bullet_id)
            p.paragraph_format.space_after = Pt(2 if compact_appendix else 4)
            p.paragraph_format.line_spacing = 1.1 if compact_appendix else 1.208
            add_inline(
                p,
                line[2:].strip(),
                font,
                east_asia,
                size=10 if compact_appendix else None,
            )
            idx += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            p = document.add_paragraph()
            apply_numbering(p, decimal_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline(p, numbered.group(1), font, east_asia)
            idx += 1
            continue

        if line.startswith("> "):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_BLUE)
            p_pr.append(shd)
            add_inline(p, line[2:].strip(), font, east_asia)
            idx += 1
            continue

        para_lines = [line]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if (
                not nxt
                or nxt.startswith(("## ", "### ", "- ", "> ", "|", "!["))
                or re.match(r"^\d+\.\s+", nxt)
            ):
                break
            para_lines.append(nxt)
            idx += 1
        prepend_explainer_space = explainer_paragraphs_remaining == 1
        p = document.add_paragraph()
        if compact_references:
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.1
        if explainer_paragraphs_remaining:
            explainer_paragraphs_remaining -= 1
            if explainer_paragraphs_remaining == 1:
                # Let the short concluding paragraph lead the next page. This
                # also keeps the first numbered heading away from the foot of
                # the executive-summary page in Word and LibreOffice.
                p.paragraph_format.space_after = Pt(72)
        if prepend_explainer_space:
            # Keep the whitespace in the paragraph that moves to the next
            # page. LibreOffice may otherwise strand a separate spacer at the
            # bottom of the previous page and place this text above the header.
            spacer_run = p.add_run()
            set_run_font(spacer_run, font, east_asia, size=11)
            for _ in range(3):
                spacer_run.add_break()
        add_inline(
            p,
            " ".join(para_lines),
            font,
            east_asia,
            size=9.5 if compact_references else None,
        )


PINGFANG_PATH = Path(
    "/System/Library/AssetsV2/com_apple_MobileAsset_Font8/86ba2c91f017a3749571a82f2c6d890ac7ffb2fb.asset/AssetData/PingFang.ttc"
)
ARIAL_PATH = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
ARIAL_BOLD_PATH = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")


def image_font(language: str, size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if language == "zh":
        return ImageFont.truetype(str(PINGFANG_PATH), size=size)
    path = ARIAL_BOLD_PATH if bold and ARIAL_BOLD_PATH.exists() else ARIAL_PATH
    return ImageFont.truetype(str(path), size=size)


def wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.FreeTypeFont,
    max_width: int,
    language: str,
) -> str:
    wrapped: list[str] = []
    for explicit_line in text.split("\n"):
        tokens = list(explicit_line) if language == "zh" else explicit_line.split(" ")
        separator = "" if language == "zh" else " "
        current = ""
        for token in tokens:
            candidate = token if not current else current + separator + token
            width = draw.textbbox((0, 0), candidate, font=font)[2]
            if current and width > max_width:
                wrapped.append(current)
                current = token
            else:
                current = candidate
        wrapped.append(current)
    return "\n".join(wrapped)


def draw_centered_box(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    text: str,
    fill: str,
    edge: str,
    language: str,
    size: int = 30,
    bold: bool = False,
    radius: int = 24,
) -> None:
    x0, y0, x1, y1 = rect
    draw.rounded_rectangle(rect, radius=radius, fill="#" + fill, outline="#" + edge, width=3)
    font = image_font(language, size, bold=bold)
    wrapped = wrap_text(draw, text, font, max_width=(x1 - x0) - 36, language=language)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=8, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((x0 + x1 - width) / 2, (y0 + y1 - height) / 2 - bbox[1]),
        wrapped,
        font=font,
        fill="#" + INK,
        spacing=8,
        align="center",
    )


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#6B7C8C", width=5)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 12), (x - 18, y + 12)], fill="#6B7C8C")


def make_logic_figure(language: str, output: Path) -> None:
    canvas = Image.new("RGB", (2100, 1050), "white")
    draw = ImageDraw.Draw(canvas)
    if language == "zh":
        title = "只改变一个事实：晚期异常是否存在"
        top = [
            "完整证据\n保留晚期异常",
            "两次独立回答",
            "时间 / 历史主题",
            "具体操作不一致\n刺激强度 vs 刺激间隔",
        ]
        bottom = [
            "反事实证据\n删除晚期异常",
            "两次独立回答",
            "空间形状主题",
            "同类操作\n固定总量改变半径",
        ]
        gates = [
            ("证据锚定", "通过", LIGHT_GREEN),
            ("方向随异常改变", "通过", LIGHT_GREEN),
            ("相同证据重复稳定", "未通过", LIGHT_RED),
        ]
        result = "按预定规则停止：不执行下游实验"
    else:
        title = "Change one fact only: whether the late anomaly exists"
        top = [
            "Full evidence\nlate anomaly retained",
            "Two independent\nresponses",
            "Temporal / history\ntheme",
            "Different operations\namplitude vs pulse lag",
        ]
        bottom = [
            "Counterfactual evidence\nlate anomaly removed",
            "Two independent\nresponses",
            "Spatial geometry\ntheme",
            "Same operation family\nradius at fixed dose",
        ]
        gates = [
            ("Evidence anchoring", "PASS", LIGHT_GREEN),
            ("Direction changes", "PASS", LIGHT_GREEN),
            ("Replicate stability", "FAIL", LIGHT_RED),
        ]
        result = "Frozen stop rule: no downstream experiment"

    title_font = image_font(language, 50, bold=True)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(
        ((2100 - (title_box[2] - title_box[0])) / 2, 44), title, font=title_font, fill="#" + INK
    )
    xs = [60, 530, 980, 1460]
    widths = [390, 330, 390, 570]
    fills = [LIGHT_BLUE, "F4F6F9", LIGHT_AMBER, LIGHT_RED]
    for row_y, texts in ((200, top), (480, bottom)):
        for i, (x, w, text) in enumerate(zip(xs, widths, texts, strict=True)):
            draw_centered_box(
                draw,
                (x, row_y, x + w, row_y + 180),
                text,
                fills[i] if i != 3 or row_y < 300 else LIGHT_GREEN,
                "7A99B6",
                language,
                size=28,
                bold=i in (0, 2),
            )
            if i < 3:
                draw_arrow(draw, (x + w + 14, row_y + 90), (xs[i + 1] - 14, row_y + 90))
    for i, (label, state, fill) in enumerate(gates):
        x = 145 + i * 650
        draw_centered_box(
            draw,
            (x, 745, x + 510, 875),
            f"{label}\n{state}",
            fill,
            "5A8A68" if fill == LIGHT_GREEN else "B45A5A",
            language,
            size=28,
            bold=True,
        )
    result_font = image_font(language, 36, bold=True)
    result_bbox = draw.textbbox((0, 0), result, font=result_font)
    draw.text(
        ((2100 - (result_bbox[2] - result_bbox[0])) / 2, 940),
        result,
        font=result_font,
        fill="#8E2F2F",
    )
    canvas.save(output, dpi=(180, 180))


def make_result_figure(language: str, output: Path) -> None:
    if language == "zh":
        columns = ["证据 / 重复", "操作", "对象", "摘要", "时刻"]
        rows = [
            ["完整 1", "第二次刺激强度", "B", "平均值", "60"],
            ["完整 2", "刺激间隔", "B", "平均值", "60"],
            ["反事实 1", "固定总量的半径", "B", "空间边界强度", "60"],
            ["反事实 2", "固定总量的半径", "B", "空间边界强度", "48"],
        ]
        title = "四次回答的操作签名"
        notes = ["强异常引用：通过", "反事实改变：通过", "完整证据重复：未通过"]
    else:
        columns = ["Evidence / run", "Intervention", "Field", "Feature", "Time"]
        rows = [
            ["Full 1", "Second-pulse amplitude", "B", "Mean", "60"],
            ["Full 2", "Pulse lag", "B", "Mean", "60"],
            ["Counterfactual 1", "Radius at fixed dose", "B", "Boundary strength", "60"],
            ["Counterfactual 2", "Radius at fixed dose", "B", "Boundary strength", "48"],
        ]
        title = "Operational signatures across four responses"
        notes = [
            "Strong-anomaly citation: PASS",
            "Counterfactual change: PASS",
            "Full-evidence stability: FAIL",
        ]

    canvas = Image.new("RGB", (2100, 950), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = image_font(language, 48, bold=True)
    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(
        ((2100 - (title_bbox[2] - title_bbox[0])) / 2, 35), title, font=title_font, fill="#" + INK
    )
    x0, y0 = 60, 135
    widths = [370, 660, 180, 560, 210]
    row_h = 118
    data = [columns, *rows]
    for r_idx, row in enumerate(data):
        x = x0
        fill = "DCE8F2" if r_idx == 0 else (LIGHT_BLUE if r_idx <= 2 else LIGHT_GREEN)
        for value, width in zip(row, widths, strict=True):
            rect = (x, y0 + r_idx * row_h, x + width, y0 + (r_idx + 1) * row_h)
            draw.rectangle(rect, fill="#" + fill, outline="#C8D2DB", width=3)
            font = image_font(language, 26 if language == "zh" else 25, bold=r_idx == 0)
            wrapped = wrap_text(draw, value, font, max_width=width - 28, language=language)
            bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6, align="center")
            tx = x + (width - (bbox[2] - bbox[0])) / 2
            ty = y0 + r_idx * row_h + (row_h - (bbox[3] - bbox[1])) / 2 - bbox[1]
            draw.multiline_text(
                (tx, ty), wrapped, font=font, fill="#" + INK, spacing=6, align="center"
            )
            x += width
    note_colors = [LIGHT_GREEN, LIGHT_GREEN, LIGHT_RED]
    edge_colors = ["#5A8A68", "#5A8A68", "#B45A5A"]
    for idx, note in enumerate(notes):
        x = 85 + idx * 680
        draw_centered_box(
            draw,
            (x, 770, x + 570, 890),
            note,
            note_colors[idx],
            edge_colors[idx].lstrip("#"),
            language,
            size=27,
            bold=True,
        )
    canvas.save(output, dpi=(180, 180))


def build(language: str, source_name: str, output_name: str) -> Path:
    document = Document()
    logic = ASSET_DIR / f"paper_logic_{language}.png"
    result = ASSET_DIR / f"paper_result_{language}.png"
    make_logic_figure(language, logic)
    make_result_figure(language, result)
    render_markdown(
        document,
        PAPER_DIR / source_name,
        language,
        {logic.name: logic, result.name: result},
    )
    output = OUTPUT_DIR / output_name
    document.save(output)
    return output


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    zh = build(
        "zh",
        "evidence_triggered_research_directions_zh.md",
        "异常是否真的改变了LLM的科研想法_中文.docx",
    )
    en = build(
        "en",
        "evidence_triggered_research_directions_en.md",
        "Does_an_Anomaly_Change_an_LLM_Research_Idea_EN.docx",
    )
    print(zh)
    print(en)


if __name__ == "__main__":
    main()
